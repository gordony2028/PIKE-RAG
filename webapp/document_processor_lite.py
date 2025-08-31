"""
Lightweight Document Processing Module for Railway Deployment
Uses OpenAI embeddings instead of sentence-transformers to stay under 4GB
"""

import os
import io
import hashlib
from typing import List, Dict, Any, Optional
from datetime import datetime
import uuid

# Document processing imports (lightweight)
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import chromadb
    from chromadb.config import Settings
except ImportError:
    chromadb = None

class DocumentProcessorLite:
    """Lightweight document processor using OpenAI embeddings"""
    
    def __init__(self, upload_dir: str, knowledge_base_dir: str, openai_client=None):
        self.upload_dir = upload_dir
        self.knowledge_base_dir = knowledge_base_dir
        self.openai_client = openai_client
        self.chroma_client = None
        
        # Ensure directories exist
        os.makedirs(upload_dir, exist_ok=True)
        os.makedirs(knowledge_base_dir, exist_ok=True)
        
        self.initialize_chroma()
    
    def initialize_chroma(self):
        """Initialize ChromaDB client with lightweight config"""
        if not chromadb:
            print("⚠️  ChromaDB not available - document indexing disabled")
            return
            
        try:
            self.chroma_client = chromadb.PersistentClient(
                path=os.path.join(self.knowledge_base_dir, "chroma_db"),
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
            print("✅ ChromaDB initialized (lite mode)")
        except Exception as e:
            print(f"⚠️  ChromaDB initialization failed: {e}")
    
    def get_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Get embeddings using OpenAI instead of sentence-transformers"""
        if not self.openai_client:
            print("⚠️  No OpenAI client - using dummy embeddings")
            return [[0.0] * 1536 for _ in texts]  # Dummy 1536-dim embeddings
        
        try:
            embeddings = []
            for text in texts:
                response = self.openai_client.embeddings.create(
                    model="text-embedding-ada-002",
                    input=text[:8000]  # Truncate to max length
                )
                embeddings.append(response.data[0].embedding)
            return embeddings
        except Exception as e:
            print(f"⚠️  Embedding generation failed: {e}")
            return [[0.0] * 1536 for _ in texts]  # Fallback dummy embeddings
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text from uploaded files"""
        try:
            if file_type == 'txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_type == 'pdf' and PyPDF2:
                text = ""
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\\n"
                return text
            
            elif file_type == 'docx':
                if not DocxDocument:
                    return "Error extracting text: python-docx library not available"
                
                try:
                    doc = DocxDocument(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():  # Skip empty paragraphs
                            text += paragraph.text + "\\n"
                    
                    # If no text was extracted, try tables
                    if not text.strip():
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip():
                                        text += cell.text + " "
                                text += "\\n"
                    
                    return text if text.strip() else "No readable text found in document"
                
                except Exception as docx_error:
                    # If it fails, it might be an older .doc file saved as .docx
                    return f"Error reading Word document: {str(docx_error)}. Note: Only .docx format (Office 2007+) is supported. Older .doc files need to be converted to .docx format."
            
            elif file_type == 'doc':
                # Old .doc format is not supported by python-docx
                return "Error: .doc format (Office 97-2003) is not supported. Please save the document as .docx format (Office 2007+) and try again."
            
            else:
                return f"Unsupported file type: {file_type}"
                
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Simple text chunking"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to end at a sentence boundary
            if end < len(text):
                # Look for sentence endings
                for i in range(end, max(start + chunk_size - 100, start), -1):
                    if text[i] in '.!?\\n':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process a document and add to knowledge base"""
        try:
            # Determine file type
            file_type = filename.split('.')[-1].lower()
            
            # Extract text
            text = self.extract_text_from_file(file_path, file_type)
            if not text or text.startswith("Error") or text.startswith("Unsupported"):
                return {
                    'success': False,
                    'error': text,
                    'document_id': None
                }
            
            # Create chunks
            chunks = self.chunk_text(text)
            
            if not self.chroma_client:
                return {
                    'success': True,
                    'message': 'Document processed (indexing disabled)',
                    'document_id': str(uuid.uuid4()),
                    'chunks': len(chunks),
                    'text_preview': text[:200] + "..." if len(text) > 200 else text
                }
            
            # Create collection
            collection_name = f"documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            collection = self.chroma_client.get_or_create_collection(
                name=collection_name,
                metadata={"filename": filename, "processed_at": datetime.now().isoformat()}
            )
            
            # Get embeddings
            embeddings = self.get_embeddings(chunks)
            
            # Add to ChromaDB
            ids = [f"chunk_{i}_{filename}" for i in range(len(chunks))]
            metadatas = [
                {
                    "filename": filename,
                    "chunk_index": i,
                    "file_type": file_type,
                    "processed_at": datetime.now().isoformat()
                }
                for i in range(len(chunks))
            ]
            
            collection.add(
                documents=chunks,
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids
            )
            
            return {
                'success': True,
                'document_id': collection_name,
                'chunks': len(chunks),
                'collection_name': collection_name,
                'text_preview': text[:200] + "..." if len(text) > 200 else text
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'document_id': None
            }
    
    def search_documents(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        """Search documents using the query"""
        if not self.chroma_client:
            return [{
                'content': 'Document search not available - ChromaDB not initialized',
                'metadata': {'note': 'Fallback response'},
                'score': 1.0
            }]
        
        try:
            # Get query embedding
            query_embedding = self.get_embeddings([query])[0]
            
            # Get all collections
            collections = self.chroma_client.list_collections()
            all_results = []
            
            for collection_info in collections:
                collection = self.chroma_client.get_collection(collection_info.name)
                
                results = collection.query(
                    query_embeddings=[query_embedding],
                    n_results=min(n_results, collection.count()),
                    include=['documents', 'metadatas', 'distances']
                )
                
                if results['documents'] and results['documents'][0]:
                    for doc, meta, dist in zip(
                        results['documents'][0],
                        results['metadatas'][0],
                        results['distances'][0]
                    ):
                        all_results.append({
                            'content': doc,
                            'metadata': meta,
                            'score': 1 - dist  # Convert distance to similarity
                        })
            
            # Sort by score and return top results
            all_results.sort(key=lambda x: x['score'], reverse=True)
            return all_results[:n_results]
            
        except Exception as e:
            print(f"Search error: {e}")
            return [{
                'content': f'Search error: {str(e)}',
                'metadata': {'error': True},
                'score': 0.0
            }]
    
    def get_available_documents(self) -> List[Dict[str, Any]]:
        """Get list of available documents"""
        if not self.chroma_client:
            return []
        
        try:
            collections = self.chroma_client.list_collections()
            documents = []
            
            for collection_info in collections:
                collection = self.chroma_client.get_collection(collection_info.name)
                metadata = collection.metadata or {}
                
                documents.append({
                    'id': collection_info.name,
                    'name': metadata.get('filename', collection_info.name),
                    'processed_at': metadata.get('processed_at', 'Unknown'),
                    'chunks': collection.count()
                })
            
            return documents
        except Exception as e:
            print(f"Error getting documents: {e}")
            return []